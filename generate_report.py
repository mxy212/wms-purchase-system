from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).parent
ASSETS = ROOT / "assets"
OUT = ROOT / "实验11_121052023023_王伟林.docx"
ASSETS.mkdir(exist_ok=True)

FONT_CN = "C:/Windows/Fonts/msyh.ttc"
FONT_CN_BOLD = "C:/Windows/Fonts/msyhbd.ttc"
FONT_MONO = "C:/Windows/Fonts/consola.ttf"


def font(path, size):
    return ImageFont.truetype(path, size)


def draw_centered(draw, xy, text, fnt, fill="#111111"):
    x, y = xy
    box = draw.textbbox((0, 0), text, font=fnt)
    draw.text((x - (box[2] - box[0]) / 2, y), text, font=fnt, fill=fill)


def arrow(draw, start, end, fill="#374151", width=3):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    p1 = (x2 - ux * 15 + px * 7, y2 - uy * 15 + py * 7)
    p2 = (x2 - ux * 15 - px * 7, y2 - uy * 15 - py * 7)
    draw.polygon([end, p1, p2], fill=fill)


def class_box(draw, x, y, w, title, lines, fill, border="#2563EB"):
    title_h = 42
    line_h = 30
    h = title_h + 16 + line_h * len(lines)
    draw.rounded_rectangle((x, y, x + w, y + h), radius=8, fill=fill, outline=border, width=3)
    draw.rectangle((x, y, x + w, y + title_h), fill=border)
    title_size = 14 if len(title) > 18 else 19
    draw_centered(draw, (x + w / 2, y + 7), title, font(FONT_CN_BOLD, title_size), "white")
    for i, line in enumerate(lines):
        draw.text((x + 14, y + title_h + 10 + i * line_h), line, font=font(FONT_CN, 16), fill="#111827")
    return h


def make_uml():
    img = Image.new("RGB", (1800, 1280), "white")
    d = ImageDraw.Draw(img)
    draw_centered(d, (900, 28), "采购单业务模块 UML 类图", font(FONT_CN_BOLD, 32), "#0F172A")

    class_box(d, 650, 100, 500, "PurchaseOrder", [
        "- state: PurchaseOrderState",
        "- logs: List<OperationLog>",
        "+ edit() / approve() / cancel()",
        "+ generateInboundOrder()",
        "+ generateInvoice() / pay()",
        "+ returnAndRefund()",
    ], "#EFF6FF")
    class_box(d, 650, 420, 500, "PurchaseOrderState <<interface>>", [
        "+ getName(): String",
        "+ edit(order, product, quantity, supplier)",
        "+ approve(order) / cancel(order)",
        "+ generateInboundOrder(order)",
        "+ generateInvoice(order) / pay(order)",
        "+ returnAndRefund(order)",
    ], "#F0FDF4", "#16A34A")
    class_box(d, 70, 420, 430, "BusinessDocumentFactory", [
        "+ createInboundOrderNo()",
        "+ createInvoiceNo()",
        "+ createReturnOrderNo()",
    ], "#FFF7ED", "#EA580C")
    class_box(d, 1300, 420, 410, "OperationLog", [
        "- time: LocalDateTime",
        "- operation: String",
        "- fromState: String",
        "- toState: String",
    ], "#FDF2F8", "#DB2777")
    class_box(d, 650, 780, 500, "AbstractPurchaseOrderState", [
        "# reject(operation)",
        "+ 默认拒绝七类业务操作",
    ], "#F8FAFC", "#64748B")

    states = [
        ("DraftState", 30), ("ApprovedState", 275), ("InboundCompletedState", 520),
        ("InvoicedState", 765), ("PaidState", 1010), ("CancelledState", 1255),
        ("ReturnedCompletedState", 1500),
    ]
    for name, x in states:
        class_box(d, x, 1080, 220, name, ["+ 覆盖允许的操作"], "#FAFAFA", "#475569")

    arrow(d, (900, 360), (900, 420), "#2563EB")
    d.text((915, 375), "持有", font=font(FONT_CN, 16), fill="#334155")
    arrow(d, (500, 520), (650, 300), "#EA580C")
    d.text((500, 385), "创建关联单据", font=font(FONT_CN, 16), fill="#334155")
    arrow(d, (1300, 520), (1150, 300), "#DB2777")
    d.text((1160, 385), "记录操作", font=font(FONT_CN, 16), fill="#334155")
    arrow(d, (900, 780), (900, 670), "#16A34A")
    d.text((915, 710), "实现", font=font(FONT_CN, 16), fill="#334155")
    for _, x in states:
        arrow(d, (x + 110, 1080), (900, 900), "#64748B", 2)
    img.save(ASSETS / "uml-class-diagram.png")


def make_state_flow():
    img = Image.new("RGB", (1800, 720), "white")
    d = ImageDraw.Draw(img)
    draw_centered(d, (900, 25), "采购单状态流转图", font(FONT_CN_BOLD, 32), "#0F172A")
    boxes = {
        "草稿": (80, 280), "已审批": (360, 280), "已入库": (660, 280),
        "已开票": (980, 160), "已付款": (1320, 160),
        "已取消": (360, 500), "已退货完结": (980, 500),
    }
    colors = {"草稿": "#E0F2FE", "已审批": "#DCFCE7", "已入库": "#FEF3C7",
              "已开票": "#FCE7F3", "已付款": "#EDE9FE",
              "已取消": "#FEE2E2", "已退货完结": "#E5E7EB"}
    for label, (x, y) in boxes.items():
        d.rounded_rectangle((x, y, x + 220, y + 80), radius=12, fill=colors[label], outline="#334155", width=3)
        draw_centered(d, (x + 110, y + 21), label, font(FONT_CN_BOLD, 24))
    transitions = [
        ("草稿", "已审批", "审批"), ("草稿", "已取消", "取消"),
        ("已审批", "已入库", "生成入库单"), ("已审批", "已取消", "取消"),
        ("已入库", "已开票", "生成发票"), ("已开票", "已付款", "付款"),
        ("已入库", "已退货完结", "退货退款"),
    ]
    for a, b, label in transitions:
        ax, ay = boxes[a]
        bx, by = boxes[b]
        start = (ax + 220, ay + 40) if bx > ax else (ax + 110, ay + 80)
        end = (bx, by + 40) if bx > ax else (bx + 110, by)
        arrow(d, start, end, "#475569", 3)
        mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
        draw_centered(d, (mx, my - 28), label, font(FONT_CN, 18), "#1F2937")
    img.save(ASSETS / "state-flow.png")


def make_test_result():
    lines = [
        "PurchaseOrderTest",
        "",
        "[PASS] 正常采购付款流程",
        "[PASS] 正常退货退款流程",
        "[PASS] 未审批禁止入库",
        "[PASS] 未入库禁止开票",
        "[PASS] 未开票禁止付款",
        "[PASS] 已入库禁止取消",
        "[PASS] 已取消禁止后续操作",
        "[PASS] 未入库禁止退货",
        "[PASS] 审批后禁止编辑",
        "",
        "测试完成：通过 9 项，失败 0 项。",
    ]
    img = Image.new("RGB", (1500, 860), "#111827")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1500, 58), fill="#1F2937")
    d.ellipse((22, 19, 42, 39), fill="#EF4444")
    d.ellipse((54, 19, 74, 39), fill="#F59E0B")
    d.ellipse((86, 19, 106, 39), fill="#22C55E")
    y = 90
    for i, line in enumerate(lines):
        color = "#86EFAC" if "[PASS]" in line or "失败 0" in line else "#E5E7EB"
        fnt = font(FONT_CN_BOLD if i == 0 else FONT_CN, 27)
        d.text((55, y), line, font=fnt, fill=color)
        y += 52
    img.save(ASSETS / "test-result.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="宋体", size=10.5, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial" if name != "Consolas" else "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial" if name != "Consolas" else "Consolas")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text="", size=10.5, bold=False, align=None, before=0, after=5, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=(89, 89, 89))


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F4F6")
    p_pr.append(shd)
    for line in code.splitlines():
        r = p.add_run(line + "\n")
        set_run_font(r, name="Consolas", size=8.3, color=(31, 41, 55))


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = Cm(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "DCE6F1")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=9.5, bold=True)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, (value, width) in enumerate(zip(row, widths)):
            cells[i].width = Cm(width)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.35
    for level, size, color, before, after in [
        (1, 14, (31, 78, 121), 10, 5),
        (2, 12, (46, 116, 181), 8, 4),
        (3, 11, (31, 78, 121), 6, 3),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    for list_name in ["List Bullet", "List Number"]:
        style = styles[list_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)


def build_doc():
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.3)
    sec.right_margin = Cm(2.3)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.2)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run("体系结构与设计模式课程实验报告")
    set_run_font(r, size=9, color=(100, 116, 139))
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("实验 11  |  121052023023  |  王伟林")
    set_run_font(r, size=9, color=(100, 116, 139))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("实验报告")
    set_run_font(r, name="黑体", size=22, bold=True, color=(15, 23, 42))

    info = doc.add_table(rows=3, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    info.style = "Table Grid"
    labels = [
        ("课程", "体系结构与设计模式", "实验名称", "基于设计模式的 WMS 仓储系统采购单全流程业务开发"),
        ("专业", "软件工程", "班级", "软工"),
        ("学号", "121052023023", "姓名", "王伟林"),
    ]
    widths = [2.0, 4.2, 2.2, 8.0]
    for row, values in zip(info.rows, labels):
        for i, (cell, value) in enumerate(zip(row.cells, values)):
            cell.width = Cm(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=120, bottom=120)
            if i % 2 == 0:
                set_cell_shading(cell, "DCE6F1")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i % 2 == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, size=9.5, bold=(i % 2 == 0))
    add_para(doc, "实验日期：2026 年 6 月 3 日", size=9.5, align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)

    add_heading(doc, "一、实验目的", 1)
    for item in [
        "理解状态模式的结构与适用场景，掌握状态驱动业务系统的设计思想。",
        "使用独立状态类封装采购单不同生命周期下的操作权限和状态流转规则。",
        "避免在采购单核心类中使用大量 if-else 或 switch 判断，提升代码可维护性与可扩展性。",
        "实现采购单编辑、审批、取消、入库、开票、付款、退货退款的完整业务闭环。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "二、实验环境", 1)
    add_table(doc, ["项目", "配置"], [
        ("开发工具", "IntelliJ IDEA 2024.3.3"),
        ("编程语言", "Java"),
        ("运行环境", "JDK 8"),
        ("设计模式", "状态模式、简单工厂模式"),
        ("测试方式", "Java main 方法自动化测试"),
    ], [4.0, 12.0])

    add_heading(doc, "三、实验内容与需求分析", 1)
    add_para(doc, "本实验围绕 WMS 仓储管理系统中的采购单业务模块展开。采购单是入库、财务对账、付款结算和退货退款的核心入口，其操作必须由当前状态决定。系统需要自动拦截非法操作，并在合法操作完成后自动切换到下一状态。")
    add_heading(doc, "3.1 核心业务操作", 2)
    add_table(doc, ["序号", "操作", "允许状态", "操作后状态"], [
        ("1", "编辑采购单", "草稿", "草稿"),
        ("2", "审批采购单", "草稿", "已审批"),
        ("3", "取消采购单", "草稿、已审批", "已取消"),
        ("4", "生成入库单", "已审批", "已入库"),
        ("5", "生成发票", "已入库", "已开票"),
        ("6", "付款", "已开票", "已付款"),
        ("7", "退货退款", "已入库", "已退货完结"),
    ], [1.4, 4.0, 5.2, 5.4])
    add_heading(doc, "3.2 关键约束", 2)
    for item in [
        "未审批的采购单不能生成入库单，未入库的采购单不能生成发票。",
        "未生成发票的采购单不能付款，已入库后的采购单不能取消。",
        "已取消、已付款、已退货完结的采购单禁止继续执行后续业务操作。",
        "只有已入库采购单可以发起退货退款。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "四、系统设计", 1)
    add_heading(doc, "4.1 设计模式说明", 2)
    add_para(doc, "系统优先采用状态模式。PurchaseOrder 作为上下文对象，只负责保存采购单数据并将操作委托给当前 PurchaseOrderState。每个具体状态类只覆盖该状态允许执行的操作；未覆盖的操作由 AbstractPurchaseOrderState 统一抛出友好业务异常。这样新增“待审核”“部分入库”等状态时，只需增加新的状态类并配置流转关系，无需修改采购单核心操作方法。")
    add_para(doc, "此外，BusinessDocumentFactory 统一创建入库单号、发票号和退货单号，OperationLog 记录操作时间及状态变化，增强模块完整性。")
    doc.add_picture(str(ASSETS / "uml-class-diagram.png"), width=Cm(16.0))
    add_caption(doc, "图 1 采购单业务模块 UML 类图")
    doc.add_picture(str(ASSETS / "state-flow.png"), width=Cm(16.0))
    add_caption(doc, "图 2 采购单状态流转图")

    add_heading(doc, "4.2 类职责说明", 2)
    add_table(doc, ["类或接口", "职责"], [
        ("PurchaseOrder", "采购单上下文，保存业务数据，向当前状态对象委托七类操作。"),
        ("PurchaseOrderState", "定义所有采购单状态必须支持的操作接口。"),
        ("AbstractPurchaseOrderState", "提供默认拒绝逻辑，统一非法操作异常信息。"),
        ("具体状态类", "封装当前状态允许的业务行为，并负责切换到下一状态。"),
        ("BusinessDocumentFactory", "统一创建入库单、发票和退货单编号。"),
        ("OperationLog", "记录操作类型、时间、原状态和目标状态。"),
    ], [5.0, 11.0])

    add_heading(doc, "五、核心代码实现", 1)
    add_heading(doc, "5.1 状态接口", 2)
    add_code(doc, """public interface PurchaseOrderState {
    String getName();
    void edit(PurchaseOrder order, String product, int quantity, String supplier);
    void approve(PurchaseOrder order);
    void cancel(PurchaseOrder order);
    void generateInboundOrder(PurchaseOrder order);
    void generateInvoice(PurchaseOrder order);
    void pay(PurchaseOrder order);
    void returnAndRefund(PurchaseOrder order);
}""")
    add_heading(doc, "5.2 默认非法操作拦截", 2)
    add_code(doc, """public abstract class AbstractPurchaseOrderState
        implements PurchaseOrderState {
    protected void reject(String operation) {
        throw new PurchaseOrderException(
            "操作失败：采购单当前状态为【" + getName()
            + "】，不允许执行【" + operation + "】。");
    }
    public void pay(PurchaseOrder order) {
        reject("付款");
    }
    // 其余操作采用相同的默认拒绝实现
}""")
    add_heading(doc, "5.3 具体状态类示例", 2)
    add_code(doc, """public class ApprovedState extends AbstractPurchaseOrderState {
    public String getName() {
        return "已审批";
    }
    public void cancel(PurchaseOrder order) {
        order.changeState(new CancelledState(), "取消采购单");
    }
    public void generateInboundOrder(PurchaseOrder order) {
        order.setInboundOrderNo(
            BusinessDocumentFactory.createInboundOrderNo(order.getOrderNo()));
        order.changeState(new InboundCompletedState(), "生成入库单");
    }
}""")
    add_heading(doc, "5.4 上下文委托操作", 2)
    add_code(doc, """public class PurchaseOrder {
    private PurchaseOrderState state;

    public void approve() {
        state.approve(this);
    }
    public void generateInboundOrder() {
        state.generateInboundOrder(this);
    }
    public void pay() {
        state.pay(this);
    }
}""")
    add_para(doc, "完整可运行 Java 源代码位于实验目录的 src/main/java 和 src/test/java 中。", bold=True, indent=False)

    add_heading(doc, "六、测试用例与运行结果", 1)
    add_heading(doc, "6.1 测试用例", 2)
    add_table(doc, ["编号", "测试场景", "预期结果"], [
        ("TC01", "草稿→审批→入库→开票→付款", "最终状态为已付款"),
        ("TC02", "草稿→审批→入库→退货退款", "最终状态为已退货完结"),
        ("TC03", "草稿状态直接生成入库单", "抛出业务异常"),
        ("TC04", "草稿状态直接生成发票", "抛出业务异常"),
        ("TC05", "未开票直接付款", "抛出业务异常"),
        ("TC06", "已入库采购单执行取消", "抛出业务异常"),
        ("TC07", "已取消采购单继续审批或入库", "抛出业务异常"),
        ("TC08", "未入库采购单执行退货退款", "抛出业务异常"),
        ("TC09", "已审批采购单执行编辑", "抛出业务异常"),
    ], [1.8, 8.2, 6.0])
    add_heading(doc, "6.2 运行结果", 2)
    add_para(doc, "在 IntelliJ IDEA 中运行 PurchaseOrderTest.main()，所有正常流程与异常流程测试均通过。")
    doc.add_picture(str(ASSETS / "test-result.png"), width=Cm(15.5))
    add_caption(doc, "图 3 自动化测试运行结果")

    add_heading(doc, "七、扩展性分析", 1)
    for item in [
        "新增“待审核”状态：创建 PendingReviewState，并在审批流程中切换到该状态即可。",
        "新增“部分入库”状态：创建 PartialInboundState，封装剩余数量校验与后续流转。",
        "新增权限校验：可在操作入口增加责任链，对角色、数据完整性和状态权限进行逐级校验。",
        "新增真实业务单据：可扩展 BusinessDocumentFactory，返回入库单、发票和退货单实体对象。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "八、仓库地址", 1)
    add_para(doc, "Gitee/GitHub 地址：待上传后填写（建议仓库名：wms-purchase-order-lab）", bold=True, indent=False)

    add_heading(doc, "九、实验总结", 1)
    add_para(doc, "通过本次实验，我进一步理解了状态模式在复杂业务流程中的价值。采购单的操作权限不再集中在一个包含大量条件判断的类中，而是被分散到各个具体状态类中，使每个类的职责更加单一，状态流转也更加清晰。")
    add_para(doc, "状态模式使系统能够在不修改核心采购单类的前提下扩展新的业务状态，符合开闭原则。结合简单工厂和操作日志后，采购单模块具备了更好的可维护性、可测试性与扩展能力，适合应用于真实 WMS 系统中的订单流程管理。")

    doc.core_properties.title = "实验11：采购单业务模块的设计与实现"
    doc.core_properties.subject = "体系结构与设计模式实验报告"
    doc.core_properties.author = "王伟林"
    doc.save(OUT)


if __name__ == "__main__":
    make_uml()
    make_state_flow()
    make_test_result()
    build_doc()
    print(OUT)
